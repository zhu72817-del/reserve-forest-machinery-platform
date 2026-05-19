from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "deliverables/储备林工程机械服务平台功能整理说明书.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(31, 77, 120)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.8)
    set_cell_text(table.rows[0].cells[0], "项目", True)
    set_cell_text(table.rows[0].cells[1], "说明", True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    style_table(table)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        if widths:
            table.columns[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    style_table(table)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 8),
        ("Subtitle", 11, "555555", 0, 12),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("储备林工程机械服务平台功能整理说明书")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(117, 117, 117)


def build():
    doc = Document()
    configure_doc(doc)
    add_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("储备林工程机械服务平台功能整理说明书")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("供功能确认、修改和后续开发调整使用")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("版本：V1.0    日期：2026年5月18日")

    doc.add_paragraph(
        "本文档依据当前平台原型和已确认的产品方向整理，重点说明平台定位、三大主模块、用户角色、前后台页面、业务流程、后台管理、数据对象和后续需确认事项。"
        "你可以直接在本 Word 文档中修改功能名称、增删模块、调整流程或补充字段；我会按你改后的文档继续调整系统。"
    )

    doc.add_heading("一、平台总体定位", level=1)
    add_kv_table(doc, [
        ("平台名称", "储备林工程机械服务平台"),
        ("建设定位", "面向储备林建设的内部采购商城、机械设备资源目录和机械服务撮合平台。"),
        ("参考模式", "参考政府采购云平台、政采云电子卖场和企业内部采购商城的业务逻辑。"),
        ("核心目标", "把内采商品、机械设备、机械服务、项目需求、询价报价、订单合同和供应商资源集中到线上管理。"),
        ("当前重点", "首页不再堆砌功能，前台围绕内采商城、机械设备、机械服务三大主模块组织。"),
    ])

    doc.add_heading("二、平台用户角色", level=1)
    add_matrix(doc, ["角色", "主要入口", "核心功能"], [
        ("游客", "前台门户", "查看首页、三大模块、项目需求大厅、供应商库、公告指南；不能下单、报价或管理数据。"),
        ("需求方 / 储备林公司", "用户中心 / 需求方工作台", "发布项目需求，浏览商品、设备和服务，发起询价，查看报价，生成订单，跟踪合同和验收。"),
        ("供应商", "供应商中心", "申请入驻，维护企业信息，上架商品/设备/服务，响应需求报价，处理订单和履约资料。"),
        ("物资公司管理员", "管理后台", "审核供应商、审核资源上架、管理需求、监管报价、订单、合同、公告和基础数据。"),
        ("超级管理员", "系统管理后台", "管理用户、角色权限、系统参数、数据字典、日志和平台配置。"),
    ], [1.2, 1.6, 3.7])

    doc.add_heading("三、首页与前台导航", level=1)
    doc.add_paragraph("首页保持简单清晰，只突出三大业务入口，避免把商城、装备、服务、供应商和项目功能混在一起。")
    add_matrix(doc, ["区域", "当前设计", "说明"], [
        ("顶部导航", "首页、内采商城、机械设备、机械服务、项目需求、供应商库", "后台管理和登录/注册作为右侧按钮，不占用主导航过多位置。"),
        ("主 Banner", "平台标题、简短定位文案、三个核心按钮", "三个按钮分别进入内采商城、机械设备、机械服务。"),
        ("三大入口卡片", "内采商城、机械设备、机械服务", "每个入口进入独立模块页面，不再只是页面锚点跳转。"),
        ("推荐内容", "推荐内采商品、推荐机械设备、推荐机械服务", "每类展示少量代表性内容，保持首页简洁。"),
        ("项目需求", "展示最新需求", "支持进入项目需求大厅或登录后发布需求。"),
        ("平台数据", "商品、设备、服务、供应商、订单、询价等数量", "数据由后台数据库统计。"),
        ("公告政策", "平台公告、操作指南、入驻流程、采购说明", "由后台内容管理维护。"),
    ], [1.3, 2.5, 2.7])

    doc.add_heading("四、三大主模块", level=1)
    doc.add_heading("4.1 内采商城", level=2)
    doc.add_paragraph("内采商城用于集团内部商品采购和租赁，主要承载森林防火装备、车辆装备、无人机、应急物资、安全防护物资和配件耗材。")
    add_matrix(doc, ["功能", "说明"], [
        ("商品分类", "森林防火装备、车辆装备、无人机装备、应急救援物资、安全防护物资、通信照明设备、配件耗材等。"),
        ("商品列表", "展示商品名称、分类、供应商、价格、区域、标签和状态。"),
        ("商品详情", "展示产品图片、型号、参数、适用场景、供货企业、价格说明、交付周期和售后服务。"),
        ("交易操作", "支持购买、租赁、发起询价、加入需求清单、联系供应商。"),
        ("后台维护", "管理员或供应商可维护商品信息、图片、参数、价格、上下架、审核和推荐状态。"),
    ], [1.5, 5.0])
    doc.add_paragraph("当前预置商品示例：")
    add_bullets(doc, [
        "森林防火装备：风力灭火机、森林消防水泵、高压细水雾灭火机、油锯、森林防火服等。",
        "车辆装备：森林消防水罐车、森林防火巡护车、指挥巡逻车、运兵车、装备运输车等。",
        "无人机装备：大疆 Matrice 350 RTK、大疆 Matrice 30T 热成像无人机、巡查无人机、电池与充电箱等。",
        "应急与耗材：应急照明灯、消防水带、无人机电池、机械滤芯、润滑油、液压油等。",
    ])

    doc.add_heading("4.2 机械设备", level=2)
    doc.add_paragraph("机械设备模块用于展示储备林工程建设、林业作业和林区运输所需设备。每个设备都应同时支持购买和租赁两个方向。")
    add_matrix(doc, ["设备分类", "示例内容"], [
        ("土地整理设备", "履带式挖掘机、小型挖掘机、轮式挖掘机、装载机、推土机、平地机、压路机。"),
        ("林地清理设备", "割灌机、林地清理机、枝丫粉碎机、树木粉碎机、油锯、清杂设备。"),
        ("造林营林设备", "植树机、挖坑机、施肥机、喷药机、灌溉设备、苗木运输设备。"),
        ("采伐集材设备", "采伐机、打枝造材机、绞盘集材机、轨道式集材系统、木材装载机。"),
        ("运输设备", "履带运输车、山地运输车、自卸运输车、木材运输车、随车吊、叉车。"),
        ("无人机设备", "大型运输无人机、巡查无人机、测绘无人机、热成像无人机。"),
    ], [1.6, 4.9])
    add_bullets(doc, [
        "购买信息：购买价格、供货周期、付款方式、质保期限、售后服务和采购说明。",
        "租赁信息：小时价、日租价、台班价、最短租期、是否含操作手、是否含运输、押金、可租赁区域和预计到场时间。",
        "详情页操作：购买、租赁、发起询价、加入需求清单、联系供应商。",
    ])

    doc.add_heading("4.3 机械服务", level=2)
    doc.add_paragraph("机械服务模块不是卖设备，而是卖工程作业服务。服务详情页必须突出服务区域、计价方式、设备人员配置和服务流程。")
    add_matrix(doc, ["服务分类", "示例内容"], [
        ("林地整理服务", "机械清杂、林地平整、作业道修整、坡地整理、机械开沟、土壤改良。"),
        ("机械化造林服务", "造林整地、挖坑、苗木栽植、补植补造、灌溉管护。"),
        ("林木抚育服务", "割灌除草、修枝抚育、施肥作业、病虫害防治、林下清理。"),
        ("采伐运输服务", "机械采伐、木材集材、木材装车、木材短驳、木材运输、山地运输。"),
        ("轨道式集材服务", "轨道安装、轨道集材作业、山地木材运输、轨道设备租赁、轨道维护。"),
        ("无人机作业服务", "无人机巡查、测绘、防火巡护、植保、物资运输、热成像监测。"),
        ("设备带人作业服务", "挖掘机带司机、装载机带司机、履带运输车带司机、割灌机作业队、采伐机械作业队。"),
    ], [1.6, 4.9])
    add_bullets(doc, [
        "价格字段：每小时工程价格、每台班价格、每亩价格、按项目报价说明。",
        "服务字段：服务供应商、服务区域、服务方式、最低服务时长、设备配置、人员配置、预计进场时间、服务保障。",
        "操作按钮：立即预约、发起询价、加入需求清单、关联设备租赁。",
    ])

    doc.add_heading("五、项目需求、询价报价、订单合同", level=1)
    add_matrix(doc, ["模块", "主要功能", "当前要求"], [
        ("项目需求", "需求方发布机械服务、设备采购、设备租赁、无人机服务、防火装备采购等需求。", "需求需进入需求大厅，支持管理员审核和供应商响应。"),
        ("询价报价", "需求方对商品、设备或服务发起询价；供应商提交报价和服务方案。", "报价可用于比选，后续可扩展揭盲报价。"),
        ("订单管理", "根据购买、租赁、服务或成交报价生成订单。", "订单类型包括内采商品购买、内采商品租赁、设备购买、设备租赁、服务订单。"),
        ("合同管理", "订单确认后生成合同，支持合同模板、附件上传、合同状态和归档。", "预留电子签章接口。"),
        ("履约验收", "供应商推进履约进度，需求方上传验收资料并评价。", "形成可追溯的订单、合同、验收和评价记录。"),
    ], [1.2, 2.7, 2.6])

    doc.add_heading("六、后台管理功能", level=1)
    add_matrix(doc, ["后台菜单", "功能范围"], [
        ("工作台", "数据概览、待办事项、快捷入口、交易趋势。"),
        ("内采商城管理", "商品列表、新增商品、商品分类、图片、参数、购买价格、租赁价格、审核、上架、下架、推荐。"),
        ("机械设备管理", "设备列表、新增设备、设备分类、设备参数、购买价格、租赁价格、设备状态、所在地区、供应商、审核、上下架。"),
        ("机械服务管理", "服务列表、新增服务、服务分类、服务区域、每小时工程价格、台班价、亩价、项目报价、服务案例、审核、上下架。"),
        ("供应商管理", "供应商列表、入驻审核、分类、标签、资质、评价和黑名单。"),
        ("项目需求管理", "需求列表、发布需求、需求审核、需求匹配、邀请供应商、需求归档。"),
        ("询价报价管理", "询价单、报价单、报价对比、揭盲报价、成交管理。"),
        ("订单管理", "购买订单、租赁订单、服务订单、履约管理、验收管理。"),
        ("合同管理", "合同列表、合同模板、合同生成、合同附件、合同归档。"),
        ("内容管理", "首页 Banner、首页三大模块、公告管理、政策文件、操作指南。"),
        ("系统管理", "用户管理、角色权限、数据字典、文件管理、系统设置、操作日志。"),
    ], [1.5, 5.0])

    doc.add_heading("七、数据库与核心数据对象", level=1)
    doc.add_paragraph("当前设计建议把三大主模块的数据分开管理，避免所有内容混成一个商品表。")
    add_matrix(doc, ["数据表 / 对象", "用途"], [
        ("mall_product", "内采商城商品，存放商品名称、分类、品牌、型号、供应商、图片、购买价、租赁价、库存、交付周期、参数和上下架状态。"),
        ("machine_equipment", "机械设备，存放设备名称、分类、品牌、型号、地区、供应商、购买价、租赁价、设备状态、作业能力和参数。"),
        ("machine_service", "机械服务，存放服务名称、分类、服务区域、小时价、台班价、亩价、项目报价说明、设备配置、人员配置和服务流程。"),
        ("supplier", "供应商信息，支持同时关联商品、设备和服务。"),
        ("project_demand", "项目需求，记录需求方、地区、类型、预算、工期、附件、状态和匹配结果。"),
        ("inquiry / quotation", "询价与报价，记录询价对象、报价金额、报价方案、附件和成交结果。"),
        ("order_main / order_item", "统一订单，按订单类型区分商品购买、商品租赁、设备购买、设备租赁和服务订单。"),
        ("contract", "合同信息，记录合同类型、模板、金额、双方主体、附件、状态和归档信息。"),
        ("file_info / audit_log", "文件附件和操作日志，用于资质、报价、合同、验收资料上传和行为留痕。"),
    ], [1.8, 4.7])

    doc.add_heading("八、主要业务流程", level=1)
    doc.add_heading("8.1 需求方采购流程", level=2)
    add_numbered(doc, [
        "需求方登录平台，进入用户中心。",
        "浏览内采商城、机械设备或机械服务，也可以直接发布项目需求。",
        "选择商品、设备或服务后，发起购买、租赁、预约或询价。",
        "供应商响应报价，需求方进行报价比选。",
        "确认供应商后生成订单，订单确认后生成合同。",
        "供应商履约，需求方验收并评价，资料归档。",
    ])
    doc.add_heading("8.2 供应商入驻与上架流程", level=2)
    add_numbered(doc, [
        "供应商提交入驻申请，填写企业资料、服务能力、资质和联系人信息。",
        "管理员审核供应商资料，通过后开通供应商权限。",
        "供应商维护内采商品、机械设备或机械服务资源。",
        "管理员审核资源，通过后在前台展示。",
        "供应商接收询价、提交报价、处理订单和履约。",
    ])
    doc.add_heading("8.3 管理员监管流程", level=2)
    add_numbered(doc, [
        "管理员进入后台工作台查看待办。",
        "审核供应商、商品、设备、服务和项目需求。",
        "监管询价报价、订单合同和履约状态。",
        "维护首页内容、公告政策、分类字典和系统配置。",
        "查看操作日志和数据统计，为后续采购决策提供支撑。",
    ])

    doc.add_heading("九、当前已实现的可演示功能", level=1)
    add_bullets(doc, [
        "前台首页已经围绕三大模块重新组织：内采商城、机械设备、机械服务。",
        "顶部导航、Banner 按钮、首页三大入口均可以进入对应独立模块页面。",
        "项目需求大厅和供应商库也可以从导航进入独立页面。",
        "购买、租赁、询价、发布需求、供应商入驻等入口已经具备可点击流程，未登录时进入登录或入驻弹窗。",
        "后端具备基础数据库、用户登录、角色区分、供应商入驻、资源上架、需求、报价、订单和日志接口。",
        "系统预置了内采商品、机械设备、机械服务、供应商、项目需求和公告数据。",
    ])

    doc.add_heading("十、需要你确认或修改的事项", level=1)
    doc.add_paragraph("建议你重点修改以下内容，改完发回给我后，我按你的确认版继续调整平台：")
    add_bullets(doc, [
        "三大模块名称是否固定为：内采商城、机械设备、机械服务。",
        "内采商城是否只放用户提供的产品资料，还是也保留部分演示通用商品。",
        "机械设备与内采商城中的设备是否需要严格分开，还是允许同一设备在两个模块同时出现。",
        "机械服务的计价字段是否固定为：每小时工程价格、台班价、亩价、项目报价。",
        "首页是否还需要保留项目需求、供应商库、平台数据和公告政策。",
        "后台菜单是否按当前三大模块重新拆分，还是需要更像政府采购云后台。",
        "是否需要增加真实商品详情页字段，例如产品参数、宣传册图片、车辆型号、技术参数表和附件下载。",
        "是否需要做公开访问的网址部署，供外部人员登录查看。",
    ])

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run("修改方式建议：").bold = True
    closing.add_run(" 你可以直接在 Word 中改文字、删除模块、增加功能点或用批注说明。我收到修改版后，会按修改内容继续开发平台。")

    doc.save(OUT)


if __name__ == "__main__":
    build()
